import json
import re

from django.conf import settings
from django.db import transaction
from django.utils import timezone

from .models import Assessment, AuditEvent, Competency, ContentImport, LearningAssignment, LearningResource, Notification, PracticeQuestion, Question, UserProfile
from .notifications import notify
from .resource_index import index_learning_resource
from .transcription import transcribe_video
from .llm.base import LLMRequest, LLMUnavailable
from .llm.factory import get_llm_provider


class ContentImportError(ValueError):
    pass


def _decode_ai_json(raw):
    """Decode the first JSON object/array even when a local model adds prose or fences."""
    cleaned = raw.strip()
    if cleaned.startswith("```"):
        cleaned = re.sub(r"^```(?:json)?\s*|\s*```$", "", cleaned, flags=re.IGNORECASE)
    decoder = json.JSONDecoder()
    for index, character in enumerate(cleaned):
        if character not in "[{":
            continue
        try:
            payload, _ = decoder.raw_decode(cleaned[index:])
        except json.JSONDecodeError:
            continue
        if isinstance(payload, (dict, list)):
            return payload
    raise ContentImportError("The AI-generated learning quiz could not be validated.")


def _balanced_source_excerpt(text, limit=8000):
    if len(text) <= limit:
        return text
    first = limit * 3 // 8
    middle = limit * 5 // 16
    last = limit - first - middle
    midpoint = len(text) // 2
    middle_start = max(first, midpoint - middle // 2)
    return f"[BEGINNING]\n{text[:first]}\n\n[MIDDLE]\n{text[middle_start:middle_start + middle]}\n\n[END]\n{text[-last:]}"


def _resolve_quiz_answer(answer, options):
    candidate = str(answer).strip()
    if candidate in options:
        return candidate
    normalized = candidate.casefold()
    exact = next((option for option in options if option.casefold() == normalized), None)
    if exact:
        return exact
    letter = re.match(r"^(?:option\s*)?([a-d])(?:[.)\s:\-]|$)", candidate, re.IGNORECASE)
    if letter:
        return options[ord(letter.group(1).upper()) - ord("A")]
    if candidate.isdigit() and 1 <= int(candidate) <= len(options):
        return options[int(candidate) - 1]
    without_label = lambda value: re.sub(r"^(?:option\s*)?[a-d][.)\s:\-]+", "", value, flags=re.IGNORECASE).strip().casefold()
    return next((option for option in options if without_label(option) == without_label(candidate)), "")


def generate_grounded_practice_questions(text, competency, target=10):
    """Generate reviewable questions grounded only in extracted material text."""
    count = min(target, max(3, len(text) // 400))
    source = _balanced_source_excerpt(text)
    response = get_llm_provider().generate(LLMRequest(
        system="You create grounded learning checks. Use only the supplied source, never invent facts, and return valid JSON only.",
        messages=[{"role": "user", "content": f'''Create exactly {count} age-appropriate multiple-choice questions from this source for {competency.code}: {competency.title}. Use four plausible options, one exact correct answer copied from the options, and a short explanation. Cover the supplied beginning, middle, and end when present. Avoid duplicate questions. Return only {{"questions":[{{"prompt":"...","options":["..."],"correct_answer":"...","explanation":"...","source_locator":"Transcript"}}]}}.\n\nSOURCE:\n{source}'''}],
        temperature=0.15, max_tokens=2200,
    ))
    payload = _decode_ai_json(response.text)
    rows = payload.get("questions", payload.get("Questions", [])) if isinstance(payload, dict) else payload
    if not isinstance(rows, list):
        raise ContentImportError("The AI-generated learning quiz did not contain a question list.")
    questions = []
    seen_prompts = set()
    for row in rows:
        if not isinstance(row, dict):
            continue
        raw_options = row.get("options", row.get("choices", []))
        if isinstance(raw_options, dict):
            raw_options = list(raw_options.values())
        options = [str(item).strip() for item in raw_options if str(item).strip()] if isinstance(raw_options, list) else []
        prompt = str(row.get("prompt", row.get("question", row.get("question_text", "")))).strip()
        raw_answer = row.get("correct_answer", row.get("correctAnswer", row.get("answer", row.get("correct", ""))))
        if not raw_answer and isinstance(row.get("answer_index"), int) and len(options) == 4 and 0 <= row["answer_index"] < 4:
            raw_answer = options[row["answer_index"]]
        answer = _resolve_quiz_answer(raw_answer, options) if len(options) == 4 else ""
        if not prompt or len(options) != 4 or not answer:
            continue
        normalized_prompt = re.sub(r"\W+", " ", prompt.casefold()).strip()
        if normalized_prompt in seen_prompts:
            continue
        seen_prompts.add(normalized_prompt)
        questions.append({"source_number": len(questions) + 1, "prompt": prompt, "question_type": Question.QuestionType.MULTIPLE_CHOICE, "options": options, "correct_answer": answer, "explanation": str(row.get("explanation", "")).strip(), "competency_id": competency.id, "competency_code": competency.code, "confidence": "needs_review", "provenance": "ai", "source_locator": str(row.get("source_locator", "Transcript"))[:240]})
        if len(questions) == count:
            break
    if len(questions) < min(3, count):
        raise ContentImportError("The AI did not return enough complete, grounded quiz questions.")
    return questions, response


def _extract_pdf(file_obj):
    from pypdf import PdfReader

    file_obj.seek(0)
    reader = PdfReader(file_obj)
    return "\n".join(page.extract_text() or "" for page in reader.pages)


def _extract_docx(file_obj):
    from docx import Document

    file_obj.seek(0)
    document = Document(file_obj)
    paragraphs = [paragraph.text for paragraph in document.paragraphs if paragraph.text.strip()]
    for table in document.tables:
        for row in table.rows:
            paragraphs.append(" | ".join(cell.text.strip() for cell in row.cells))
    return "\n".join(paragraphs)


def extract_document_text(content_import):
    filename = content_import.original_filename.casefold()
    if filename.endswith(".pdf"):
        text = _extract_pdf(content_import.source_file)
    elif filename.endswith(".docx"):
        text = _extract_docx(content_import.source_file)
    else:
        raise ContentImportError("Only text-based PDF and DOCX documents can be extracted.")
    text = text.replace("\x00", "").strip()
    if len(text) < 20:
        raise ContentImportError("No usable text was found. Scanned documents require OCR and cannot be imported in this release.")
    return text


QUESTION_START = re.compile(r"(?m)^\s*(\d+)[.)]\s+")
CHOICE_LINE = re.compile(r"^\s*([A-Fa-f])[.)]\s*(.+?)\s*$")
INLINE_CHOICE_MARKER = re.compile(r"(?<!\w)([A-Fa-f])[.)]\s*")
ANSWER_LINE = re.compile(r"^\s*(?:answer|correct answer|key)\s*:\s*(.+?)\s*$", re.IGNORECASE)
COMPETENCY_LINE = re.compile(r"^\s*competenc(?:y|ies)\s*:\s*(.+?)\s*$", re.IGNORECASE)
ANSWER_KEY_ITEM = re.compile(r"^\s*(\d+)\s*[.)]\s*(.+?)\s*$")
MATCHING_CHOICE_LINE = re.compile(r"^\s*([A-Z])[.)]\s*(.+?)\s*$")


def _choices_from_line(line):
    matches = list(INLINE_CHOICE_MARKER.finditer(line))
    if not matches or line[:matches[0].start()].strip():
        return []
    return [{"label": match.group(1).upper(), "text": line[match.end(): matches[index + 1].start() if index + 1 < len(matches) else len(line)].strip()} for index, match in enumerate(matches) if line[match.end(): matches[index + 1].start() if index + 1 < len(matches) else len(line)].strip()]


def parse_exam_text(text, default_competency=None, answer_key=None):
    starts = list(QUESTION_START.finditer(text))
    questions = []
    for index, match in enumerate(starts):
        block = text[match.end(): starts[index + 1].start() if index + 1 < len(starts) else len(text)]
        lines = [line.strip() for line in block.splitlines() if line.strip()]
        prompt_lines, options, answer, competency_code = [], [], "", ""
        for line in lines:
            inline_options = _choices_from_line(line)
            choice_match = CHOICE_LINE.match(line)
            answer_match = ANSWER_LINE.match(line)
            competency_match = COMPETENCY_LINE.match(line)
            if inline_options:
                options.extend(inline_options)
            elif choice_match:
                options.append({"label": choice_match.group(1).upper(), "text": choice_match.group(2)})
            elif answer_match:
                answer = answer_match.group(1).strip()
            elif competency_match:
                competency_code = competency_match.group(1).strip()
            elif not options and not answer and not re.fullmatch(r"\d{1,3}", line):
                prompt_lines.append(line)
        prompt = " ".join(prompt_lines).strip()
        answer = answer or str((answer_key or {}).get(int(match.group(1)), "")).strip()
        if not prompt or not answer:
            continue
        if len(answer) == 1 and options:
            selected = next((item["text"] for item in options if item["label"] == answer.upper()), None)
            answer = selected or answer
        competency = default_competency
        if competency_code:
            competency = Competency.objects.filter(subject=default_competency.subject if default_competency else None, code__iexact=competency_code).first()
        question_type = Question.QuestionType.MULTIPLE_CHOICE if len(options) > 2 else Question.QuestionType.TRUE_FALSE if len(options) == 2 and {item["text"].casefold() for item in options} == {"true", "false"} else Question.QuestionType.SHORT_ANSWER
        questions.append({
            "source_number": int(match.group(1)),
            "prompt": prompt,
            "question_type": question_type,
            "options": [item["text"] for item in options],
            "correct_answer": answer,
            "competency_id": competency.id if competency else None,
            "competency_code": competency.code if competency else competency_code,
            "confidence": "high" if competency and (options or question_type == Question.QuestionType.SHORT_ANSWER) else "needs_review",
        })
    return questions


def _module_section(lines, heading, start=0):
    normalized_heading = heading.casefold()
    candidates = [index for index, line in enumerate(lines[start:], start=start) if line.strip().rstrip(":").casefold() == normalized_heading]
    return candidates[-1] if candidates else None


def _parse_matching_assessment(section_lines, answer_key, competency):
    """Convert a Column A/Column B matching task into reviewable four-choice items."""
    if "column a" not in " ".join(section_lines).casefold() or "column b" not in " ".join(section_lines).casefold():
        return []
    questions, choices = [], {}
    current_question = None
    current_choice = None
    reading_choices = False
    for raw_line in section_lines:
        line = raw_line.strip()
        if not line or line.casefold() in {"column a", "column b", "column a column b"} or re.fullmatch(r"\d{1,3}", line):
            continue
        choice_match = MATCHING_CHOICE_LINE.match(line)
        question_match = QUESTION_START.match(line)
        if choice_match and questions:
            reading_choices = True
            current_question = None
            current_choice = choice_match.group(1)
            choices[current_choice] = choice_match.group(2).strip()
        elif question_match and not reading_choices:
            current_question = {"number": int(question_match.group(1)), "prompt": line[question_match.end():].strip()}
            questions.append(current_question)
        elif reading_choices and current_choice:
            choices[current_choice] = f"{choices[current_choice]} {line}".strip()
        elif current_question:
            current_question["prompt"] = f"{current_question['prompt']} {line}".strip()
    if len(questions) < 2 or len(choices) < 4:
        return []
    choice_labels = list(choices)
    parsed = []
    for item in questions:
        correct_label = str(answer_key.get(item["number"], "")).strip().upper()
        correct_answer = choices.get(correct_label)
        if not correct_answer:
            continue
        correct_index = choice_labels.index(correct_label)
        distractors = [choices[choice_labels[(correct_index + offset) % len(choice_labels)]] for offset in range(1, len(choice_labels))][:3]
        options = distractors[:]
        options.insert((item["number"] - 1) % 4, correct_answer)
        parsed.append({
            "source_number": item["number"],
            "prompt": item["prompt"],
            "question_type": Question.QuestionType.MULTIPLE_CHOICE,
            "options": options,
            "correct_answer": correct_answer,
            "competency_id": competency.id if competency else None,
            "competency_code": competency.code if competency else "",
            "confidence": "needs_review",
            "provenance": "extracted",
            "source_locator": f"Assessment matching item {item['number']}",
        })
    return parsed


def parse_module_questions(text, default_competency=None):
    """Extract the end-of-module assessment and its separately printed answer key."""
    lines = text.splitlines()
    answer_key_start = _module_section(lines, "Answer Key")
    if answer_key_start is None:
        return parse_exam_text(text, default_competency)
    assessment_start = _module_section(lines[:answer_key_start], "Assessment")
    section_name = "Assessment"
    if assessment_start is None:
        assessment_start = _module_section(lines[:answer_key_start], "What I Know")
        section_name = "What I Know"
    if assessment_start is None:
        return parse_exam_text(text, default_competency)
    section_end = next((index for index in range(assessment_start + 1, answer_key_start) if lines[index].strip().rstrip(":").casefold() in {"additional activities", "answer key", "references"}), answer_key_start)
    key_section_start = next((index for index in range(answer_key_start + 1, len(lines)) if lines[index].strip().rstrip(":").casefold() == section_name.casefold()), None)
    if key_section_start is None:
        return parse_exam_text("\n".join(lines[assessment_start + 1:section_end]), default_competency)
    answer_key = {}
    for line in lines[key_section_start + 1:]:
        stripped = line.strip()
        if stripped.endswith(":") and not ANSWER_KEY_ITEM.match(stripped):
            break
        match = ANSWER_KEY_ITEM.match(stripped)
        if match:
            answer_key[int(match.group(1))] = match.group(2).strip()
    section_lines = lines[assessment_start + 1:section_end]
    matching_questions = _parse_matching_assessment(section_lines, answer_key, default_competency)
    return matching_questions or parse_exam_text("\n".join(section_lines), default_competency, answer_key)


def process_content_import(content_import):
    content_import.status = ContentImport.Status.PROCESSING
    content_import.error_message = ""
    content_import.save(update_fields=["status", "error_message", "updated_at"])
    try:
        if content_import.kind == ContentImport.Kind.VIDEO:
            if settings.WHISPER_ENABLED:
                text = transcribe_video(content_import)
                practice_questions = parse_module_questions(text, content_import.competency)
                generation = {"quiz_generation_status": "extracted" if practice_questions else "not_attempted"}
                if not practice_questions and content_import.competency:
                    try:
                        practice_questions, generated = generate_grounded_practice_questions(text, content_import.competency)
                        generation = {"quiz_generation_status": "ai_generated", "quiz_provider": generated.provider, "quiz_model": generated.model}
                    except (LLMUnavailable, ContentImportError, ValueError) as exc:
                        generation = {"quiz_generation_status": "unavailable", "quiz_generation_error": str(exc)}
                payload = {"resource_type": LearningResource.ResourceType.VIDEO, "description": content_import.configuration.get("description", ""), "transcription_status": "completed", "transcript_character_count": len(text), "practice_questions": practice_questions, "practice_question_count": len(practice_questions), **generation}
            else:
                payload = {"resource_type": LearningResource.ResourceType.VIDEO, "description": content_import.configuration.get("description", ""), "transcription_status": "not_configured", "practice_questions": [], "practice_question_count": 0}
                text = ""
        else:
            text = extract_document_text(content_import)
            if content_import.kind == ContentImport.Kind.EXAM:
                questions = parse_exam_text(text, content_import.competency)
                if not questions:
                    raise ContentImportError("No complete questions were detected. Use numbered questions and include an 'Answer:' line for each item.")
                payload = {"questions": questions, "question_count": len(questions)}
            else:
                practice_questions = parse_module_questions(text, content_import.competency)
                payload = {"resource_type": LearningResource.ResourceType.MODULE, "character_count": len(text), "practice_questions": practice_questions, "practice_question_count": len(practice_questions)}
        content_import.extracted_text = text
        content_import.extracted_payload = payload
        content_import.status = ContentImport.Status.NEEDS_REVIEW
    except Exception as exc:
        content_import.status = ContentImport.Status.FAILED
        content_import.error_message = str(exc)[:2000]
    content_import.save(update_fields=["extracted_text", "extracted_payload", "status", "error_message", "updated_at"])
    return content_import


def _validate_exam_payload(content_import):
    questions = content_import.extracted_payload.get("questions", [])
    if not questions:
        raise ContentImportError("At least one extracted question is required.")
    validated = []
    for position, item in enumerate(questions, start=1):
        competency = Competency.objects.filter(pk=item.get("competency_id"), subject=content_import.subject).first()
        if not competency:
            raise ContentImportError(f"Question {position} must be mapped to a valid competency.")
        prompt = str(item.get("prompt", "")).strip()
        correct_answer = str(item.get("correct_answer", "")).strip()
        question_type = item.get("question_type")
        options = [str(option).strip() for option in item.get("options", []) if str(option).strip()]
        if not prompt or not correct_answer or question_type not in Question.QuestionType.values:
            raise ContentImportError(f"Question {position} is incomplete.")
        if question_type in {Question.QuestionType.MULTIPLE_CHOICE, Question.QuestionType.TRUE_FALSE} and correct_answer not in options:
            raise ContentImportError(f"Question {position}'s correct answer must match one of its choices.")
        validated.append((competency, prompt, question_type, options, correct_answer))
    return validated


@transaction.atomic
def sync_published_practice_questions(content_import, actor):
    if content_import.kind not in {ContentImport.Kind.MODULE, ContentImport.Kind.VIDEO} or not content_import.published_resource_id:
        return 0
    questions = content_import.extracted_payload.get("practice_questions", [])
    validated = []
    for position, item in enumerate(questions, start=1):
        prompt = str(item.get("prompt", "")).strip()
        answer = str(item.get("correct_answer", "")).strip()
        question_type = item.get("question_type")
        options = [str(option).strip() for option in item.get("options", []) if str(option).strip()]
        if not prompt or not answer or question_type not in PracticeQuestion.QuestionType.values:
            raise ContentImportError(f"Quiz question {position} is incomplete.")
        if question_type in {PracticeQuestion.QuestionType.MULTIPLE_CHOICE, PracticeQuestion.QuestionType.TRUE_FALSE} and answer not in options:
            raise ContentImportError(f"Quiz question {position}'s correct answer must match one of its choices.")
        validated.append((prompt, question_type, options, answer, str(item.get("explanation", "")), str(item.get("provenance", "extracted")), str(item.get("source_locator", ""))))
    resource = content_import.published_resource
    resource.practice_questions.all().delete()
    PracticeQuestion.objects.bulk_create([
        PracticeQuestion(resource=resource, prompt=prompt, question_type=question_type, options=options, correct_answer=answer, explanation=explanation, provenance=provenance, source_locator=locator, position=position)
        for position, (prompt, question_type, options, answer, explanation, provenance, locator) in enumerate(validated, start=1)
    ])
    AuditEvent.objects.create(actor=actor, action="content_import.quiz_revised", object_type="ContentImport", object_id=str(content_import.id), metadata={"question_count": len(validated), "resource_id": resource.id})
    return len(validated)


@transaction.atomic
def publish_content_import(content_import, reviewer):
    if content_import.status != ContentImport.Status.NEEDS_REVIEW:
        raise ContentImportError("Only imports awaiting review can be published.")
    if content_import.kind == ContentImport.Kind.EXAM:
        questions = _validate_exam_payload(content_import)
        assessment = Assessment.objects.create(
            title=content_import.title,
            subject=content_import.subject,
            kind=content_import.configuration.get("assessment_kind", Assessment.Kind.PRE),
            instructions=content_import.configuration.get("instructions", ""),
            due_at=content_import.configuration.get("due_at") or None,
            is_active=False,
            created_by=reviewer,
        )
        class_ids = content_import.configuration.get("assigned_class_ids", [])
        assessment.assigned_classes.set(class_ids)
        for competency, prompt, question_type, options, answer in questions:
            Question.objects.create(assessment=assessment, competency=competency, prompt=prompt, question_type=question_type, options=options, correct_answer=answer)
        content_import.published_assessment = assessment
    else:
        resource_type = LearningResource.ResourceType.VIDEO if content_import.kind == ContentImport.Kind.VIDEO else LearningResource.ResourceType.MODULE
        resource = LearningResource.objects.create(
            title=content_import.title,
            resource_type=resource_type,
            difficulty=content_import.configuration.get("difficulty", "foundation"),
            content=(f"{content_import.configuration.get('description', '').strip()}\n\nTranscript:\n{content_import.extracted_text}".strip() if resource_type == LearningResource.ResourceType.VIDEO else content_import.extracted_text),
            file=content_import.source_file.name,
            original_filename=content_import.original_filename,
            mime_type=content_import.mime_type,
            is_approved=True,
            uploaded_by=content_import.uploaded_by,
            purpose=content_import.configuration.get("purpose", LearningResource.Purpose.REGULAR),
        )
        if content_import.competency_id:
            resource.competencies.add(content_import.competency)
        if resource_type in {LearningResource.ResourceType.MODULE, LearningResource.ResourceType.VIDEO}:
            for position, item in enumerate(content_import.extracted_payload.get("practice_questions", []), start=1):
                prompt = str(item.get("prompt", "")).strip()
                answer = str(item.get("correct_answer", "")).strip()
                question_type = item.get("question_type")
                options = [str(option).strip() for option in item.get("options", []) if str(option).strip()]
                if prompt and answer and question_type in PracticeQuestion.QuestionType.values:
                    PracticeQuestion.objects.create(resource=resource, prompt=prompt, question_type=question_type, options=options, correct_answer=answer, explanation=str(item.get("explanation", "")), provenance=str(item.get("provenance", "extracted")), source_locator=str(item.get("source_locator", ""))[:240], position=position)
        index_learning_resource(resource)
        content_import.published_resource = resource
        class_ids = content_import.configuration.get("assigned_class_ids", [])
        if class_ids:
            assignment = LearningAssignment.objects.create(resource=resource, assigned_by=content_import.uploaded_by, instructions=content_import.configuration.get("instructions", ""), due_at=content_import.configuration.get("due_at") or None)
            assignment.assigned_classes.set(class_ids)
            students = UserProfile.objects.filter(role=UserProfile.Role.STUDENT, academic_class_id__in=class_ids, is_active=True).select_related("user").distinct()
            for profile in students:
                notify(recipient=profile.user, kind=Notification.Kind.LEARNING_ASSIGNED, title="Learning material assigned", message=f"{resource.title} has been assigned to your class.", action_url="/materials", deduplication_key=f"learning-assignment:{assignment.id}:student:{profile.user_id}")
    content_import.status = ContentImport.Status.PUBLISHED
    content_import.reviewed_by = reviewer
    content_import.reviewed_at = timezone.now()
    content_import.save(update_fields=["published_assessment", "published_resource", "status", "reviewed_by", "reviewed_at", "updated_at"])
    AuditEvent.objects.create(actor=reviewer, action="content_import.published", object_type="ContentImport", object_id=str(content_import.pk), metadata={"kind": content_import.kind, "title": content_import.title})
    return content_import
