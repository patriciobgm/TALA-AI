import json
import tempfile
from unittest.mock import patch

from django.contrib.auth import get_user_model
from django.core import signing
from django.core.cache import cache
from django.core.management import call_command
from django.core.files.uploadedfile import SimpleUploadedFile
from django.test import override_settings
from django.utils import timezone
from docx import Document
from io import BytesIO
from urllib.parse import urlsplit
from rest_framework.test import APITestCase
from .account_security import _totp
from .content_imports import _extract_pdf, generate_grounded_practice_questions, parse_module_questions
from .learning_intelligence import rank_learning_resources
from .llm.base import LLMResponse
from .models import AIMessage, AcademicClass, Assessment, AssessmentAttempt, AssessmentEligibility, Competency, CompetencyResult, ContentImport, EnrollmentRequest, LearnerCompetencyEvidence, LearningAssignment, LearningResource, LearningResourceChunk, Notification, PrivacyAcknowledgment, Question, RecoveryActivity, RecoveryPlan, RemedialExamConsent, ResearchEvaluationSnapshot, StudentProfile, Subject, UserProfile
from .services import create_recovery_plan

class RecoveryApiTests(APITestCase):
    @classmethod
    def setUpTestData(cls):
        call_command("seed_demo", verbosity=0)

    def setUp(self):
        cache.clear()

    def authenticate(self, username="student@tala.edu.ph", password="demo-password"):
        response = self.client.post("/api/auth/login/", {"username": username, "password": password}, format="json")
        self.assertEqual(response.status_code, 200)
        self.client.credentials(HTTP_AUTHORIZATION=f"Bearer {response.data['access']}")
        return response.data

    def correct_practice_answers(self, activity):
        return {str(question.id): question.correct_answer for question in activity.resource.practice_questions.all()}

    @patch("pypdf.PdfReader")
    def test_pdf_extraction_reads_reader_pages(self, reader_class):
        first_page = reader_class.return_value.pages = [reader_class.return_value.page]
        first_page[0].extract_text.return_value = "Extracted module content"

        text = _extract_pdf(BytesIO(b"pdf"))

        self.assertEqual(text, "Extracted module content")

    def test_module_parser_matches_separate_deped_answer_key_and_inline_choices(self):
        competency = Competency.objects.get(code="GM-03")
        text = """Assessment
1. Which practice protects an online account?
A. Share passwords  B. Use a strong password  C. Ignore updates  D. Post credentials
2. What should you do with a suspicious link?
A. Open it  B. Forward it  C. Avoid it  D. Save it
Additional Activities
Answer Key
Assessment:
1.B
2.C
References
"""

        questions = parse_module_questions(text, competency)

        self.assertEqual(len(questions), 2)
        self.assertEqual(questions[0]["options"], ["Share passwords", "Use a strong password", "Ignore updates", "Post credentials"])
        self.assertEqual(questions[0]["correct_answer"], "Use a strong password")
        self.assertEqual(questions[1]["correct_answer"], "Avoid it")

    def test_module_parser_converts_matching_columns_to_valid_mcq_items(self):
        competency = Competency.objects.get(code="GM-03")
        text = """Assessment
Column A  Column B
1. Find an exact phrase
2. Find education websites
A. plain words
B. \"exact phrase\"
C. Site:.edu
D. Filetype:pdf
Additional Activities
Answer Key
Assessment:
1.B
2.C
References
"""

        questions = parse_module_questions(text, competency)

        self.assertEqual(len(questions), 2)
        self.assertEqual(questions[0]["correct_answer"], '"exact phrase"')
        self.assertEqual(questions[1]["correct_answer"], "Site:.edu")
        self.assertTrue(all(len(item["options"]) == 4 for item in questions))
        self.assertTrue(all(item["correct_answer"] in item["options"] for item in questions))
        self.assertTrue(all(item["confidence"] == "needs_review" for item in questions))

    def test_login_includes_role_context(self):
        data = self.authenticate()
        self.assertEqual(data["user"]["role"], "student")
        self.assertEqual(data["user"]["class_name"], "Grade 11 – Rizal")
        self.assertTrue(data["user"]["privacy_acknowledgment_required"])

    def test_student_acknowledges_current_privacy_declaration(self):
        self.authenticate()
        current = self.client.get("/api/auth/privacy-acknowledgment/")
        self.assertEqual(current.status_code, 200, current.data)
        self.assertFalse(current.data["acknowledged"])
        accepted = self.client.post("/api/auth/privacy-acknowledgment/", {"accepted": True}, format="json")
        self.assertEqual(accepted.status_code, 200, accepted.data)
        student = get_user_model().objects.get(username="student@tala.edu.ph")
        self.assertTrue(PrivacyAcknowledgment.objects.filter(user=student, policy_version=current.data["policy_version"]).exists())
        self.assertFalse(self.client.get("/api/auth/me/").data["privacy_acknowledgment_required"])

    @override_settings(EMAIL_BACKEND="django.core.mail.backends.locmem.EmailBackend")
    def test_new_student_gets_onboarding_email_and_must_change_password(self):
        from django.core import mail
        self.authenticate("admin@tala.edu.ph")
        created = self.client.post("/api/users/", {"name": "New ARAL Learner", "email": "new.aral@tala.edu.ph", "role": "student", "password": "temporary-password-42", "academic_class": None, "is_active": True}, format="json")
        self.assertEqual(created.status_code, 201, created.data)
        self.assertTrue(created.data["must_change_password"])
        self.assertEqual(created.data["assignment"], "Unassigned")
        self.assertEqual(len(mail.outbox), 1)
        self.assertIn("set your password", mail.outbox[0].subject.lower())
        self.assertIn("/reset-password?", mail.outbox[0].body)

    def test_student_class_code_request_requires_teacher_approval(self):
        student = get_user_model().objects.get(username="student@tala.edu.ph")
        target = AcademicClass.objects.exclude(pk=student.tala_profile.academic_class_id).filter(grade_level=11).first()
        self.assertTrue(target.class_code)
        current_teacher = get_user_model().objects.filter(tala_profile__role=UserProfile.Role.TEACHER, tala_profile__assigned_classes=student.tala_profile.academic_class).first()
        self.authenticate(current_teacher.username)
        removed = self.client.post("/api/enrollment-requests/unenroll/", {"student": student.id}, format="json")
        self.assertEqual(removed.status_code, 200, removed.data)
        target_teacher = get_user_model().objects.filter(tala_profile__role=UserProfile.Role.TEACHER, tala_profile__assigned_classes=target).first()
        subject = target_teacher.tala_profile.assigned_subjects.filter(grade_level=target.grade_level).first()
        self.authenticate()
        requested = self.client.post("/api/enrollment-requests/", {"class_code": f"{subject.code}-{target.class_code}"}, format="json")
        self.assertEqual(requested.status_code, 201, requested.data)
        self.assertEqual(requested.data["status"], EnrollmentRequest.Status.PENDING)
        student.tala_profile.refresh_from_db()
        self.assertNotEqual(student.tala_profile.academic_class_id, target.id)
        teacher = get_user_model().objects.filter(tala_profile__role=UserProfile.Role.TEACHER, tala_profile__assigned_classes=target).first()
        self.assertIsNotNone(teacher)
        self.authenticate(teacher.username)
        decided = self.client.post(f"/api/enrollment-requests/{requested.data['id']}/decide/", {"decision": "approved"}, format="json")
        self.assertEqual(decided.status_code, 200, decided.data)
        student.tala_profile.refresh_from_db()
        self.assertEqual(student.tala_profile.academic_class_id, target.id)

    def test_teacher_enrollment_candidates_follow_selected_subject_grade(self):
        teacher = get_user_model().objects.get(username="teacher@tala.edu.ph")
        subject = teacher.tala_profile.assigned_subjects.order_by("id").first()
        other_grade = 12 if subject.grade_level == 11 else 11
        matching = get_user_model().objects.create_user(username="matching.candidate@example.com", email="matching.candidate@example.com", password="test-password", first_name="Matching")
        matching_profile = UserProfile.objects.create(user=matching, role=UserProfile.Role.STUDENT)
        StudentProfile.objects.create(profile=matching_profile, grade_level=subject.grade_level)
        other = get_user_model().objects.create_user(username="other.candidate@example.com", email="other.candidate@example.com", password="test-password", first_name="Other")
        other_profile = UserProfile.objects.create(user=other, role=UserProfile.Role.STUDENT)
        StudentProfile.objects.create(profile=other_profile, grade_level=other_grade)
        self.authenticate(teacher.username)

        response = self.client.get(f"/api/enrollment-requests/candidates/?mode=enroll&subject={subject.id}")

        self.assertEqual(response.status_code, 200, response.data)
        candidate_ids = {item["id"] for item in response.data}
        self.assertIn(matching.id, candidate_ids)
        self.assertNotIn(other.id, candidate_ids)
        self.assertTrue(all(item["grade_level"] == subject.grade_level for item in response.data))

    def test_teacher_must_de_enroll_before_enrolling_learner_again(self):
        student = get_user_model().objects.get(username="student@tala.edu.ph")
        teacher = get_user_model().objects.filter(tala_profile__role=UserProfile.Role.TEACHER, tala_profile__assigned_classes=student.tala_profile.academic_class).first()
        current_class = student.tala_profile.academic_class
        self.authenticate(teacher.username)

        subject = teacher.tala_profile.assigned_subjects.filter(grade_level=current_class.grade_level).first()
        blocked = self.client.post("/api/enrollment-requests/", {"student": student.id, "academic_class": current_class.id, "subject": subject.id}, format="json")
        self.assertEqual(blocked.status_code, 409, blocked.data)
        self.assertIn("Unenroll", blocked.data["detail"])

        removed = self.client.post("/api/enrollment-requests/unenroll/", {"student": student.id}, format="json")
        self.assertEqual(removed.status_code, 200, removed.data)
        student.tala_profile.refresh_from_db()
        self.assertIsNone(student.tala_profile.academic_class_id)

        candidates = self.client.get(f"/api/enrollment-requests/candidates/?mode=enroll&subject={subject.id}")
        self.assertIn(student.id, {item["id"] for item in candidates.data})
        enrolled = self.client.post("/api/enrollment-requests/", {"student": student.id, "academic_class": current_class.id, "subject": subject.id}, format="json")
        self.assertEqual(enrolled.status_code, 201, enrolled.data)
        student.tala_profile.refresh_from_db()
        self.assertEqual(student.tala_profile.academic_class_id, current_class.id)

    def test_notification_can_be_dismissed_only_by_recipient(self):
        self.authenticate()
        student = get_user_model().objects.get(username="student@tala.edu.ph")
        notification = Notification.objects.filter(recipient=student).first()
        dismissed = self.client.delete(f"/api/notifications/{notification.id}/dismiss/")
        self.assertEqual(dismissed.status_code, 204)
        self.assertFalse(Notification.objects.filter(pk=notification.id).exists())

    def test_privacy_request_notifies_administrators_and_status_notifies_requester(self):
        self.authenticate()
        created = self.client.post(
            "/api/privacy/requests/",
            {"request_type": "access", "details": "Please provide a copy of my profile data."},
            format="json",
        )
        self.assertEqual(created.status_code, 201, created.data)
        request_id = created.data["id"]
        administrators = get_user_model().objects.filter(tala_profile__role=UserProfile.Role.ADMIN, is_active=True)
        self.assertEqual(administrators.count(), 2)
        for administrator in administrators:
            notification = Notification.objects.get(
                recipient=administrator,
                kind=Notification.Kind.PRIVACY_REQUEST,
                deduplication_key=f"privacy-request:{request_id}:submitted:admin:{administrator.id}",
            )
            self.assertEqual(notification.action_url, f"/settings/privacy/{request_id}")
            self.assertNotIn("profile data", notification.message.lower())

        self.authenticate("admin@tala.edu.ph")
        reviewing = self.client.patch(
            "/api/privacy/requests/",
            {"id": request_id, "status": "in_review"},
            format="json",
        )
        self.assertEqual(reviewing.status_code, 200, reviewing.data)
        student = get_user_model().objects.get(username="student@tala.edu.ph")
        reviewing_notification = Notification.objects.get(
            recipient=student,
            deduplication_key=f"privacy-request:{request_id}:in_review:requester",
        )
        self.assertEqual(reviewing_notification.action_url, "/profile")

        completed = self.client.patch(
            "/api/privacy/requests/",
            {"id": request_id, "status": "completed", "resolution": "Secure export provided to the requester."},
            format="json",
        )
        self.assertEqual(completed.status_code, 200, completed.data)
        self.assertTrue(Notification.objects.filter(
            recipient=student,
            deduplication_key=f"privacy-request:{request_id}:completed:requester",
        ).exists())

    def test_demo_dataset_has_multiple_teacher_and_student_scenarios(self):
        self.assertEqual(UserProfile.objects.filter(role=UserProfile.Role.TEACHER).count(), 3)
        self.assertEqual(UserProfile.objects.filter(role=UserProfile.Role.STUDENT).count(), 10)
        self.assertTrue(RecoveryPlan.objects.filter(student__username="student@tala.edu.ph").exists())
        nina = get_user_model().objects.get(username="nina.flores@tala.edu.ph")
        self.assertFalse(Assessment.objects.filter(subject__code="GM", assigned_classes=nina.tala_profile.academic_class).exists())
        self.assertEqual(Subject.objects.count(), 5)
        self.assertEqual(Competency.objects.count(), 20)
        self.assertEqual(AcademicClass.objects.count(), 3)

    def test_superadmin_can_manage_administrator_access_but_admin_cannot(self):
        response = self.client.post("/api/auth/login/", {"username": "superadmin@tala.edu.ph", "password": "demo-password"}, format="json")
        self.assertEqual(response.status_code, 200, response.data)
        self.assertTrue(response.data["user"]["is_superadmin"])
        self.client.credentials(HTTP_AUTHORIZATION=f"Bearer {response.data['access']}")
        administrators = self.client.get("/api/users/?role=admin")
        self.assertEqual(administrators.status_code, 200)
        self.assertTrue(any(item["email"] == "admin@tala.edu.ph" for item in administrators.data["results"]))
        self.authenticate("admin@tala.edu.ph")
        hidden = self.client.get("/api/users/?role=admin")
        self.assertEqual(hidden.status_code, 200)
        self.assertEqual(hidden.data["count"], 0)

    def test_user_management_defaults_to_role_appropriate_sorting(self):
        self.authenticate("admin@tala.edu.ph")
        students = self.client.get("/api/users/?role=student&page_size=100")
        teachers = self.client.get("/api/users/?role=teacher&page_size=100")
        User = get_user_model()

        expected_students = list(User.objects.filter(tala_profile__role=UserProfile.Role.STUDENT).order_by("tala_profile__student_details__grade_level", "last_name", "first_name", "id").values_list("id", flat=True))
        expected_teachers = list(User.objects.filter(tala_profile__role=UserProfile.Role.TEACHER).order_by("last_name", "first_name", "id").values_list("id", flat=True))
        self.assertEqual([item["id"] for item in students.data["results"]], expected_students)
        self.assertEqual([item["id"] for item in teachers.data["results"]], expected_teachers)
        self.assertEqual(User.objects.get(username="teacher@tala.edu.ph").tala_profile.get_role_display(), "Teacher")

    def test_admin_can_manage_curriculum_classes_and_assignments(self):
        self.authenticate("admin@tala.edu.ph")
        subject = self.client.post("/api/subjects/", {"name": "Media and Information Literacy", "code": "MIL", "grade_level": 12, "is_active": True}, format="json")
        self.assertEqual(subject.status_code, 201, subject.data)
        grade_twelve = self.client.get("/api/subjects/?grade=12")
        self.assertEqual(grade_twelve.status_code, 200, grade_twelve.data)
        self.assertEqual([item["code"] for item in grade_twelve.data["results"]], ["MIL"])
        competency = self.client.post("/api/competencies/", {"subject": subject.data["id"], "code": "MIL-01", "title": "Evaluate information sources", "mastery_threshold": 80, "is_active": True}, format="json")
        self.assertEqual(competency.status_code, 201, competency.data)
        academic_class = self.client.post("/api/classes/", {"name": "Luna", "grade_level": 12, "school_year": "2026-2027", "is_active": True}, format="json")
        self.assertEqual(academic_class.status_code, 201, academic_class.data)
        user = self.client.post("/api/users/", {"name": "Test Learner", "email": "test.learner@tala.edu.ph", "role": "student", "password": "defensible-password-42", "academic_class": academic_class.data["id"], "is_active": True}, format="json")
        self.assertEqual(user.status_code, 201, user.data)
        updated = self.client.patch(f"/api/users/{user.data['id']}/", {"is_active": False}, format="json")
        self.assertEqual(updated.status_code, 200, updated.data)
        self.assertEqual(updated.data["status"], "Inactive")

    def test_teacher_class_access_is_derived_from_assigned_subject_grades(self):
        self.authenticate("admin@tala.edu.ph")
        teacher = get_user_model().objects.get(username="teacher@tala.edu.ph")
        grade_twelve_subject = Subject.objects.create(name="Grade Twelve Test Subject", code="G12-TEST", grade_level=12, is_active=True)
        response = self.client.patch(f"/api/users/{teacher.id}/", {"assigned_subjects": [grade_twelve_subject.id], "assigned_classes": []}, format="json")
        self.assertEqual(response.status_code, 200, response.data)
        expected = set(AcademicClass.objects.filter(grade_level=12, is_active=True).values_list("id", flat=True))
        self.assertEqual(set(response.data["assigned_classes"]), expected)
        self.assertEqual(set(teacher.tala_profile.assigned_classes.values_list("id", flat=True)), expected)

    def test_admin_can_freeze_a_checksummed_research_evidence_snapshot(self):
        self.authenticate("admin@tala.edu.ph")
        current = self.client.get("/api/research/evidence/")
        self.assertEqual(current.status_code, 200, current.data)
        self.assertEqual(current.data["algorithm_version"], "evidence-rank-v1")
        self.assertIn("diagnostic_to_mastery", current.data["metrics"])
        frozen = self.client.post("/api/research/snapshots/", {"name": "Pilot baseline", "dataset_version": "pilot-v1", "notes": "Pre-registration baseline"}, format="json")
        self.assertEqual(frozen.status_code, 201, frozen.data)
        self.assertEqual(len(frozen.data["checksum_sha256"]), 64)
        snapshot = ResearchEvaluationSnapshot.objects.get(pk=frozen.data["id"])
        self.assertEqual(snapshot.algorithm_version, "evidence-rank-v1")

    def test_teacher_can_add_question_to_draft_assessment(self):
        self.authenticate("teacher@tala.edu.ph")
        assessment = Assessment.objects.filter(subject__in=get_user_model().objects.get(username="teacher@tala.edu.ph").tala_profile.assigned_subjects.all()).first()
        assessment.is_active = False
        assessment.save(update_fields=["is_active"])
        competency = assessment.subject.competencies.first()

        response = self.client.post(
            f"/api/assessments/{assessment.id}/questions/",
            {
                "competency": competency.id,
                "prompt": "Which response demonstrates the target competency?",
                "question_type": "mcq",
                "options": ["Option A", "Option B", "Option C", "Option D"],
                "correct_answer": "Option B",
            },
            format="json",
        )

        self.assertEqual(response.status_code, 201, response.data)
        self.assertEqual(response.data["competency"], competency.id)
        self.assertTrue(assessment.questions.filter(pk=response.data["id"]).exists())
        edited = self.client.patch(f"/api/assessments/{assessment.id}/questions/{response.data['id']}/", {"prompt": "Which revised response demonstrates the target competency?"}, format="json")
        self.assertEqual(edited.status_code, 200, edited.data)
        self.assertIn("revised", edited.data["prompt"])

    def test_teacher_assessment_automatically_uses_subject_classes(self):
        self.authenticate("teacher@tala.edu.ph")
        teacher = get_user_model().objects.get(username="teacher@tala.edu.ph")
        subject = teacher.tala_profile.assigned_subjects.order_by("id").first()
        expected_classes = set(teacher.tala_profile.assigned_classes.filter(is_active=True, grade_level=subject.grade_level).values_list("id", flat=True))
        unrelated_class = AcademicClass.objects.exclude(id__in=expected_classes).first()

        response = self.client.post("/api/assessments/", {
            "title": "Automatically scoped diagnostic",
            "subject": subject.id,
            "kind": Assessment.Kind.PRE,
            "assigned_classes": [unrelated_class.id] if unrelated_class else [],
        }, format="json")

        self.assertEqual(response.status_code, 201, response.data)
        self.assertEqual(set(response.data["assigned_classes"]), expected_classes)

    def test_profile_password_and_mfa_workflow(self):
        self.authenticate("teacher@tala.edu.ph")
        profile = self.client.patch("/api/auth/profile/", {"name": "Elena Test Cruz"}, format="json")
        self.assertEqual(profile.status_code, 200, profile.data)
        setup = self.client.post("/api/auth/mfa/setup/", {}, format="json")
        self.assertEqual(setup.status_code, 200, setup.data)
        confirmed = self.client.post("/api/auth/mfa/confirm/", {"otp": _totp(setup.data["secret"])}, format="json")
        self.assertEqual(confirmed.status_code, 200, confirmed.data)
        self.assertEqual(len(confirmed.data["recovery_codes"]), 8)
        self.client.credentials()
        requires_mfa = self.client.post("/api/auth/login/", {"username": "teacher@tala.edu.ph", "password": "demo-password"}, format="json")
        self.assertEqual(requires_mfa.status_code, 400)
        signed_in = self.client.post("/api/auth/login/", {"username": "teacher@tala.edu.ph", "password": "demo-password", "otp": _totp(setup.data["secret"])}, format="json")
        self.assertEqual(signed_in.status_code, 200, signed_in.data)

    def test_student_dashboard_uses_persisted_plans(self):
        self.authenticate()
        response = self.client.get("/api/dashboard/student/")
        self.assertEqual(response.status_code, 200)
        self.assertGreaterEqual(len(response.data["plans"]), 1)
        self.assertGreaterEqual(len(response.data["attempts"]), 1)

    def test_grade_12_nina_does_not_receive_grade_11_general_mathematics(self):
        nina = get_user_model().objects.get(email="nina.flores@tala.edu.ph")
        mathematics = Subject.objects.get(code="GM")
        self.authenticate(nina.username)

        context = self.client.get("/api/dashboard/student/context/")
        assessments = self.client.get(f"/api/assessments/?subject={mathematics.id}")

        self.assertNotIn(mathematics.id, {item["id"] for item in context.data["subjects"]})
        self.assertEqual(assessments.data["results"], [])

    def test_student_dashboard_prioritizes_pending_diagnostic(self):
        student = get_user_model().objects.get(username="student@tala.edu.ph")
        diagnostic = Assessment.objects.get(kind=Assessment.Kind.PRE, subject__code="GM")
        AssessmentAttempt.objects.filter(student=student, assessment=diagnostic).delete()
        RecoveryPlan.objects.filter(student=student, competency__subject=diagnostic.subject).delete()
        self.authenticate()

        response = self.client.get(f"/api/dashboard/student/?subject={diagnostic.subject_id}")

        self.assertEqual(response.status_code, 200, response.data)
        self.assertEqual(response.data["pending_diagnostic"]["id"], diagnostic.id)
        self.assertEqual(response.data["pending_diagnostic"]["question_count"], diagnostic.questions.count())

    def test_diagnostic_waits_for_explicit_prerequisite_material(self):
        student = get_user_model().objects.get(username="student@tala.edu.ph")
        teacher = get_user_model().objects.get(username="teacher@tala.edu.ph")
        diagnostic = Assessment.objects.get(kind=Assessment.Kind.PRE, subject__code="GM")
        AssessmentAttempt.objects.filter(student=student, assessment=diagnostic).delete()
        material = LearningResource.objects.create(title="Required diagnostic preparation", resource_type=LearningResource.ResourceType.MODULE, content="Review the class introduction.", is_approved=True, uploaded_by=teacher, purpose=LearningResource.Purpose.PREREQUISITE)
        material.competencies.add(Competency.objects.get(code="GM-01"))
        assignment = LearningAssignment.objects.create(resource=material, assigned_by=teacher)
        assignment.assigned_classes.add(student.tala_profile.academic_class)
        diagnostic.prerequisite_assignments.add(assignment)
        self.authenticate()

        blocked = self.client.get(f"/api/assessments/{diagnostic.id}/start/")
        self.assertEqual(blocked.status_code, 409, blocked.data)
        completed = self.client.post(f"/api/learning-assignments/{assignment.id}/complete/", {}, format="json")
        self.assertEqual(completed.status_code, 200, completed.data)
        started = self.client.get(f"/api/assessments/{diagnostic.id}/start/")
        self.assertEqual(started.status_code, 200, started.data)

    def test_remedial_exam_is_visible_only_after_teacher_confirms_student(self):
        student = get_user_model().objects.get(username="student@tala.edu.ph")
        remedial = Assessment.objects.get(kind=Assessment.Kind.REMEDIAL)
        self.authenticate()
        before = self.client.get("/api/assessments/")
        self.assertNotIn(remedial.id, {item["id"] for item in before.data["results"]})
        self.authenticate("teacher@tala.edu.ph")
        confirmed = self.client.post(f"/api/assessments/{remedial.id}/eligibility/", {"student": student.id, "status": AssessmentEligibility.Status.ELIGIBLE, "reason": "Teacher review confirmed."}, format="json")
        self.assertEqual(confirmed.status_code, 200, confirmed.data)
        self.authenticate()
        after = self.client.get("/api/assessments/")
        self.assertIn(remedial.id, {item["id"] for item in after.data["results"]})

    def test_student_academic_workspace_is_scoped_by_subject(self):
        student = get_user_model().objects.get(username="student@tala.edu.ph")
        teacher = get_user_model().objects.get(username="teacher@tala.edu.ph")
        english = Subject.objects.get(code="EAPP")
        english_competency = english.competencies.first()
        plan, _ = RecoveryPlan.objects.get_or_create(student=student, competency=english_competency, defaults={"baseline_score": 45, "status": "active"})
        resource = LearningResource.objects.create(title="English scoped module", resource_type=LearningResource.ResourceType.MODULE, content="English lesson", is_approved=True, uploaded_by=teacher)
        resource.competencies.add(english_competency)
        assignment = LearningAssignment.objects.create(resource=resource, assigned_by=teacher)
        assignment.assigned_classes.add(student.tala_profile.academic_class)
        assessment = Assessment.objects.create(title="English scoped assessment", subject=english, kind=Assessment.Kind.PRE, is_active=True, created_by=teacher)
        assessment.assigned_classes.add(student.tala_profile.academic_class)
        self.authenticate()

        context = self.client.get("/api/dashboard/student/context/")
        dashboard = self.client.get(f"/api/dashboard/student/?subject={english.id}")
        plans = self.client.get(f"/api/recovery-plans/?subject={english.id}")
        materials = self.client.get(f"/api/learning-assignments/?subject={english.id}")
        assessments = self.client.get(f"/api/assessments/?subject={english.id}")

        self.assertEqual(context.status_code, 200, context.data)
        self.assertIn(english.id, {item["id"] for item in context.data["subjects"]})
        self.assertEqual(dashboard.data["academic_class"]["class_code"], f"{english.code}-{student.tala_profile.academic_class.class_code}")
        self.assertEqual({item["competency"] for item in dashboard.data["plans"]}, {english_competency.id})
        self.assertEqual({item["id"] for item in plans.data["results"]}, {plan.id})
        self.assertEqual({item["id"] for item in materials.data["results"]}, {assignment.id})
        self.assertIn(assessment.id, {item["id"] for item in assessments.data["results"]})
        self.assertTrue(all(item["subject"] == english.id for item in assessments.data["results"]))

    def test_activities_must_be_completed_in_order(self):
        self.authenticate()
        plan = RecoveryPlan.objects.filter(activities__resource__isnull=False).distinct().first()
        activities = list(plan.activities.order_by("position"))
        blocked = self.client.post(f"/api/recovery-plans/{plan.id}/activities/{activities[1].id}/complete/", {"answers": {}}, format="json")
        self.assertEqual(blocked.status_code, 409)
        completed = self.client.post(f"/api/recovery-plans/{plan.id}/activities/{activities[0].id}/complete/", {"answers": self.correct_practice_answers(activities[0])}, format="json")
        self.assertEqual(completed.status_code, 200)
        self.assertTrue(completed.data["passed"])
        self.assertEqual(completed.data["feedback"][0]["student_answer"], completed.data["feedback"][0]["correct_answer"])
        self.assertEqual(completed.data["activity"]["review"]["feedback"][0]["correct_answer"], completed.data["feedback"][0]["correct_answer"])

    def test_locked_mastery_assessment_explains_remaining_work(self):
        self.authenticate()
        post_assessment = Assessment.objects.get(kind=Assessment.Kind.POST)
        response = self.client.get("/api/assessments/")
        self.assertEqual(response.status_code, 200)
        serialized = next(item for item in response.data["results"] if item["id"] == post_assessment.id)
        self.assertFalse(serialized["available"])
        self.assertGreater(serialized["remaining_activities"], 0)
        self.assertIn("remaining recovery", serialized["availability_reason"])
        plans = self.client.get("/api/recovery-plans/").data["results"]
        linked_plan = next(item for item in plans if item["competency"] in post_assessment.questions.values_list("competency_id", flat=True))
        self.assertEqual(linked_plan["mastery_assessment"]["id"], post_assessment.id)
        self.assertFalse(linked_plan["mastery_assessment"]["available"])
        expected_for_competency = RecoveryPlan.objects.get(pk=linked_plan["id"]).activities.filter(resource__isnull=False, completed_at__isnull=True).count()
        self.assertEqual(linked_plan["mastery_assessment"]["remaining_activities"], expected_for_competency)
        self.assertLessEqual(linked_plan["mastery_assessment"]["remaining_activities"], serialized["remaining_activities"])

    def test_wrong_practice_answer_does_not_complete_activity(self):
        self.authenticate()
        plan = RecoveryPlan.objects.filter(activities__resource__isnull=False).distinct().first()
        activity = plan.activities.filter(resource__isnull=False).order_by("position").first()
        question = activity.resource.practice_questions.first()
        response = self.client.post(f"/api/recovery-plans/{plan.id}/activities/{activity.id}/complete/", {"answers": {str(question.id): "definitely incorrect"}}, format="json")
        self.assertEqual(response.status_code, 200)
        self.assertFalse(response.data["passed"])
        self.assertTrue(LearnerCompetencyEvidence.objects.filter(student__username="student@tala.edu.ph", competency=plan.competency, evidence_type=LearnerCompetencyEvidence.EvidenceType.PRACTICE, source_id=response.data["attempt"]["id"]).exists())
        activity.refresh_from_db()
        self.assertIsNone(activity.completed_at)

    @patch("recovery.tutor.get_llm_provider")
    def test_tutor_is_learner_aware_and_returns_chunk_citations(self, provider_factory):
        provider_factory.return_value.generate.return_value = LLMResponse(text="Use a common denominator before adding.", provider="test", model="test-model")
        self.authenticate()
        plan = RecoveryPlan.objects.get(student__username="student@tala.edu.ph", competency__code="GM-03")
        response = self.client.post(f"/api/tutor/plans/{plan.id}/messages/", {"message": "Why can I not add these directly?", "action": "explain"}, format="json")
        self.assertEqual(response.status_code, 200, response.data)
        self.assertEqual(response.data["grounding_status"], "grounded")
        self.assertGreater(len(response.data["sources"]), 0)
        self.assertIn("[1]", response.data["answer"])
        self.assertTrue(LearningResourceChunk.objects.filter(resource_id=response.data["sources"][0]["resource_id"], resource__competencies=plan.competency, resource__is_approved=True).exists())
        saved = AIMessage.objects.get(pk=response.data["id"])
        self.assertEqual(saved.source_citations, response.data["sources"])
        request = provider_factory.return_value.generate.call_args.args[0]
        self.assertIn("Learner evidence summary", request.system)
        self.assertIn("never an instruction", request.system)
        feedback = self.client.post(f"/api/tutor/messages/{response.data['id']}/feedback/", {"rating": "helpful"}, format="json")
        self.assertEqual(feedback.status_code, 200, feedback.data)
        follow_up = self.client.post(f"/api/tutor/plans/{plan.id}/messages/", {"message": "Give me a hint for the next step.", "action": "hint"}, format="json")
        self.assertEqual(follow_up.status_code, 200, follow_up.data)
        follow_up_request = provider_factory.return_value.generate.call_args.args[0]
        self.assertEqual([item["role"] for item in follow_up_request.messages], ["user", "assistant", "user"])

    def test_tutor_rejects_unknown_instruction_modes(self):
        self.authenticate()
        plan = RecoveryPlan.objects.filter(student__username="student@tala.edu.ph").first()
        response = self.client.post(f"/api/tutor/plans/{plan.id}/messages/", {"message": "Do something", "action": "change_grade"}, format="json")
        self.assertEqual(response.status_code, 400)

    @patch("recovery.tutor.get_llm_provider")
    def test_learning_quiz_tutor_is_assignment_grounded_and_does_not_disclose_answers(self, provider_factory):
        provider_factory.return_value.generate.return_value = LLMResponse(text="Review the relevant module section before choosing.", provider="test", model="test-model")
        self.authenticate("student@tala.edu.ph")
        nina = get_user_model().objects.get(username="student@tala.edu.ph")
        competency = Competency.objects.get(code="GM-03")
        resource = LearningResource.objects.create(title="Grounded quiz module", resource_type=LearningResource.ResourceType.MODULE, content="Use a common denominator before adding unlike fractions.", is_approved=True)
        resource.competencies.add(competency)
        question = resource.practice_questions.create(prompt="What should you identify before adding unlike fractions?", question_type="mcq", options=["A common denominator", "A decimal", "A product"], correct_answer="A common denominator")
        assignment = LearningAssignment.objects.create(resource=resource, assigned_by=get_user_model().objects.get(username="teacher@tala.edu.ph"))
        assignment.assigned_classes.add(nina.tala_profile.academic_class)

        response = self.client.post(
            f"/api/tutor/learning-assignments/{assignment.id}/messages/",
            {"message": "Give me a hint for this question.", "action": "hint", "question_id": question.id, "selected_answer": question.options[0]},
            format="json",
        )

        self.assertEqual(response.status_code, 200, response.data)
        self.assertEqual(response.data["grounding_status"], "grounded")
        self.assertTrue(response.data["sources"])
        llm_request = provider_factory.return_value.generate.call_args.args[0]
        self.assertIn(assignment.resource.title, llm_request.system)
        self.assertIn(question.prompt, llm_request.system)
        self.assertIn("never reveal, confirm", llm_request.system)
        saved = AIMessage.objects.get(pk=response.data["id"])
        self.assertEqual(saved.conversation.learning_assignment_id, assignment.id)
        self.assertIsNone(saved.conversation.plan_id)

    def test_student_cannot_access_teacher_dashboard(self):
        self.authenticate()
        response = self.client.get("/api/dashboard/teacher/learners/")
        self.assertEqual(response.status_code, 403)

    def test_teacher_can_access_learner_dashboard(self):
        self.authenticate("teacher@tala.edu.ph")
        response = self.client.get("/api/dashboard/teacher/learners/")
        self.assertEqual(response.status_code, 200)
        self.assertEqual(response.data[0]["name"], "Maria Santos")

    def test_teacher_dashboard_metrics_follow_selected_subject(self):
        teacher = get_user_model().objects.get(username="teacher@tala.edu.ph")
        mathematics = Subject.objects.get(code="GM")
        english = Subject.objects.get(code="EAPP")
        self.assertEqual(teacher.tala_profile.assigned_subjects.filter(pk__in=[mathematics.id, english.id]).count(), 2)
        self.authenticate("teacher@tala.edu.ph")

        math_rows = self.client.get(f"/api/dashboard/teacher/learners/?subject={mathematics.id}")
        english_rows = self.client.get(f"/api/dashboard/teacher/learners/?subject={english.id}")

        self.assertEqual(math_rows.status_code, 200, math_rows.data)
        self.assertEqual(english_rows.status_code, 200, english_rows.data)
        maria_math = next(item for item in math_rows.data if item["email"] == "student@tala.edu.ph")
        maria_english = next(item for item in english_rows.data if item["email"] == "student@tala.edu.ph")
        self.assertGreater(maria_math["gaps"], maria_english["gaps"])

    def test_teacher_workspace_collections_follow_selected_subject(self):
        teacher = get_user_model().objects.get(username="teacher@tala.edu.ph")
        mathematics = Subject.objects.get(code="GM")
        english = Subject.objects.get(code="EAPP")
        math_competency = mathematics.competencies.first()
        english_competency = english.competencies.first()
        math_resource = LearningResource.objects.create(title="Math scope resource", resource_type=LearningResource.ResourceType.LESSON, content="Math", is_approved=True, uploaded_by=teacher)
        english_resource = LearningResource.objects.create(title="English scope resource", resource_type=LearningResource.ResourceType.LESSON, content="English", is_approved=True, uploaded_by=teacher)
        math_resource.competencies.add(math_competency)
        english_resource.competencies.add(english_competency)
        import_defaults = {"mime_type": "application/pdf", "size_bytes": 4, "checksum_sha256": "test-checksum"}
        ContentImport.objects.create(title="Math scope import", kind=ContentImport.Kind.MODULE, subject=mathematics, competency=math_competency, source_file="content-imports/math.pdf", original_filename="math.pdf", uploaded_by=teacher, **import_defaults)
        ContentImport.objects.create(title="English scope import", kind=ContentImport.Kind.MODULE, subject=english, competency=english_competency, source_file="content-imports/english.pdf", original_filename="english.pdf", uploaded_by=teacher, **import_defaults)
        self.authenticate("teacher@tala.edu.ph")

        resources = self.client.get(f"/api/resources/?subject={mathematics.id}")
        imports = self.client.get(f"/api/content-imports/?subject={mathematics.id}")

        self.assertEqual(resources.status_code, 200, resources.data)
        resource_titles = {item["title"] for item in resources.data["results"]}
        self.assertIn(math_resource.title, resource_titles)
        self.assertNotIn(english_resource.title, resource_titles)
        self.assertEqual(imports.status_code, 200, imports.data)
        import_titles = {item["title"] for item in imports.data["results"]}
        self.assertIn("Math scope import", import_titles)
        self.assertNotIn("English scope import", import_titles)

    def test_teacher_can_accept_an_evidence_ranked_material_recommendation(self):
        teacher = get_user_model().objects.get(username="teacher@tala.edu.ph")
        student = get_user_model().objects.get(username="student@tala.edu.ph")
        plan = RecoveryPlan.objects.filter(student=student, status="active", competency__subject__in=teacher.tala_profile.assigned_subjects.all()).first()
        resource = LearningResource.objects.create(title="Targeted guided example", resource_type=LearningResource.ResourceType.EXAMPLE, difficulty="guided", content="A teacher-approved guided example.", is_approved=True, uploaded_by=teacher)
        resource.competencies.add(plan.competency)
        recommendation = rank_learning_resources(student, plan.competency, exclude_resource_ids=plan.activities.values_list("resource_id", flat=True), limit=10)
        self.assertIn(resource.id, {item["resource"].id for item in recommendation})
        mastery_position = plan.activities.get(resource__isnull=True).position

        self.authenticate("teacher@tala.edu.ph")
        response = self.client.post(f"/api/dashboard/teacher/learners/{student.id}/recommendations/", {"plan": plan.id, "resource": resource.id, "decision": "accepted"}, format="json")

        self.assertEqual(response.status_code, 201, response.data)
        activity = plan.activities.get(resource=resource)
        self.assertEqual(activity.position, mastery_position)
        self.assertTrue(activity.recommendation_reason)
        self.assertEqual(plan.activities.get(resource__isnull=True).position, mastery_position + 1)

    def test_teacher_can_edit_questions_in_an_assigned_draft_assessment(self):
        teacher = get_user_model().objects.get(username="teacher@tala.edu.ph")
        assessment = Assessment.objects.filter(subject__in=teacher.tala_profile.assigned_subjects.all(), questions__isnull=False).distinct().first()
        assessment.is_active = False
        assessment.save(update_fields=["is_active"])
        question = assessment.questions.first()

        self.authenticate("teacher@tala.edu.ph")
        response = self.client.patch(f"/api/assessments/{assessment.id}/questions/{question.id}/", {"prompt": "Revised teacher question"}, format="json")

        self.assertEqual(response.status_code, 200, response.data)
        question.refresh_from_db()
        self.assertEqual(question.prompt, "Revised teacher question")

    def test_teacher_learner_detail_is_limited_to_assigned_subjects(self):
        teacher = get_user_model().objects.get(username="liza.navarro@tala.edu.ph")
        student = get_user_model().objects.filter(tala_profile__academic_class__in=teacher.tala_profile.assigned_classes.all(), tala_profile__role=UserProfile.Role.STUDENT).first()
        unauthorized_competency = Competency.objects.exclude(subject__in=teacher.tala_profile.assigned_subjects.all()).first()
        LearnerCompetencyEvidence.objects.create(student=student, competency=unauthorized_competency, evidence_type=LearnerCompetencyEvidence.EvidenceType.PRACTICE, source_type="test", source_id=999, summary="Evidence outside the teacher's subject scope.")

        self.authenticate("liza.navarro@tala.edu.ph")
        response = self.client.get(f"/api/dashboard/teacher/learners/{student.id}/")

        self.assertEqual(response.status_code, 200, response.data)
        allowed_subjects = set(teacher.tala_profile.assigned_subjects.values_list("id", flat=True))
        returned_evidence_subjects = {item["subject"] for item in response.data["evidence"]}
        returned_result_subjects = {result["subject"] for attempt in response.data["attempts"] for result in attempt["competency_results"]}
        self.assertTrue(returned_evidence_subjects.issubset(allowed_subjects))
        self.assertTrue(returned_result_subjects.issubset(allowed_subjects))
        self.assertNotIn(unauthorized_competency.subject_id, returned_evidence_subjects)

    def test_teacher_can_load_content_import_dependencies(self):
        self.authenticate("teacher@tala.edu.ph")
        for path in ["/api/content-imports/", "/api/subjects/", "/api/competencies/", "/api/classes/"]:
            response = self.client.get(path)
            self.assertEqual(response.status_code, 200, path)
        subjects = self.client.get("/api/subjects/").data["results"]
        assigned_subject_ids = set(UserProfile.objects.get(user__username="teacher@tala.edu.ph").assigned_subjects.values_list("id", flat=True))
        self.assertEqual({item["id"] for item in subjects}, assigned_subject_ids)
        classes = self.client.get("/api/classes/").data["results"]
        self.assertEqual(classes[0]["label"], "Grade 11 – Rizal")
        context = self.client.get("/api/dashboard/teacher/context/")
        self.assertEqual(context.status_code, 200, context.data)
        self.assertEqual({item["id"] for item in context.data["subjects"]}, assigned_subject_ids)
        self.assertEqual({item["id"] for item in context.data["classes"]}, set(UserProfile.objects.get(user__username="teacher@tala.edu.ph").assigned_classes.values_list("id", flat=True)))

    def test_complete_recovery_and_submit_mastery_assessment(self):
        self.authenticate()
        plans_response = self.client.get("/api/recovery-plans/")
        plans = plans_response.data["results"]
        for plan in plans:
            for activity in sorted(plan["activities"], key=lambda item: item["position"]):
                if activity["resource"]:
                    model_activity = RecoveryPlan.objects.get(pk=plan["id"]).activities.get(pk=activity["id"])
                    response = self.client.post(f"/api/recovery-plans/{plan['id']}/activities/{activity['id']}/complete/", {"answers": self.correct_practice_answers(model_activity)}, format="json")
                    self.assertEqual(response.status_code, 200)
        post_assessment = Assessment.objects.get(kind=Assessment.Kind.POST)
        start = self.client.get(f"/api/assessments/{post_assessment.id}/start/")
        self.assertEqual(start.status_code, 200)
        answers = [{"question_id": question.id, "answer": question.correct_answer} for question in post_assessment.questions.all()]
        submitted = self.client.post(f"/api/assessments/{post_assessment.id}/submit/", {"answers": answers}, format="json")
        self.assertEqual(submitted.status_code, 201)
        self.assertEqual(float(submitted.data["score"]), 100)
        self.assertFalse(RecoveryPlan.objects.filter(student__username="student@tala.edu.ph", status="active").exists())
        self.assertFalse(RecoveryPlan.objects.filter(student__username="student@tala.edu.ph", activities__resource__isnull=True, activities__completed_at__isnull=True).exists())

    def test_old_diagnostic_result_does_not_reopen_mastered_plan(self):
        student = get_user_model().objects.get(username="student@tala.edu.ph")
        old_result = CompetencyResult.objects.filter(
            attempt__student=student,
            competency__code="GM-04",
            status=CompetencyResult.Status.REMEDIATION,
        ).first()
        plan = RecoveryPlan.objects.get(student=student, competency=old_result.competency)
        plan.status = "completed"
        plan.save(update_fields=["status"])
        latest_attempt = AssessmentAttempt.objects.create(
            assessment=Assessment.objects.get(kind=Assessment.Kind.POST),
            student=student,
            submitted_at=timezone.now(),
            score=100,
        )
        CompetencyResult.objects.create(
            attempt=latest_attempt,
            competency=old_result.competency,
            score=100,
            status=CompetencyResult.Status.MASTERED,
        )
        stale_plan = RecoveryPlan.objects.create(
            student=student,
            competency=old_result.competency,
            baseline_score=old_result.score,
            status="active",
        )

        returned = create_recovery_plan(student, old_result)

        self.assertEqual(returned.id, plan.id)
        self.assertFalse(RecoveryPlan.objects.filter(pk=stale_plan.id).exists())
        self.assertFalse(RecoveryPlan.objects.filter(student=student, competency=old_result.competency, status="active").exists())

    def test_mastery_check_can_target_one_recovery_competency(self):
        self.authenticate()
        plan = RecoveryPlan.objects.get(student__username="student@tala.edu.ph", competency__code="GM-03")
        for activity in plan.activities.filter(resource__isnull=False).order_by("position"):
            response = self.client.post(f"/api/recovery-plans/{plan.id}/activities/{activity.id}/complete/", {"answers": self.correct_practice_answers(activity)}, format="json")
            self.assertEqual(response.status_code, 200, response.data)
        assessment = Assessment.objects.get(kind=Assessment.Kind.POST)
        started = self.client.get(f"/api/assessments/{assessment.id}/start/?competency={plan.competency_id}")
        self.assertEqual(started.status_code, 200, started.data)
        questions = started.data["assessment"]["questions"]
        self.assertEqual(len(questions), 2)
        self.assertTrue(all(question["competency"] == plan.competency_id for question in questions))
        answers = [{"question_id": question.id, "answer": question.correct_answer} for question in assessment.questions.filter(competency=plan.competency)]
        submitted = self.client.post(f"/api/assessments/{assessment.id}/submit/", {"competency": plan.competency_id, "answers": answers}, format="json")
        self.assertEqual(submitted.status_code, 201, submitted.data)
        plan.refresh_from_db()
        self.assertEqual(plan.status, "completed")
        self.assertTrue(plan.activities.filter(resource__isnull=True, completed_at__isnull=False).exists())

    def test_mastery_retry_contains_only_questions_missed_in_latest_attempt(self):
        self.authenticate()
        plan = RecoveryPlan.objects.get(student__username="student@tala.edu.ph", competency__code="GM-03")
        for activity in plan.activities.filter(resource__isnull=False).order_by("position"):
            completed = self.client.post(f"/api/recovery-plans/{plan.id}/activities/{activity.id}/complete/", {"answers": self.correct_practice_answers(activity)}, format="json")
            self.assertEqual(completed.status_code, 200, completed.data)
        assessment = Assessment.objects.get(kind=Assessment.Kind.POST)
        questions = list(assessment.questions.filter(competency=plan.competency))
        self.assertEqual(len(questions), 2)
        first_answers = [
            {"question_id": questions[0].id, "answer": questions[0].correct_answer},
            {"question_id": questions[1].id, "answer": "deliberately incorrect"},
        ]
        submitted = self.client.post(f"/api/assessments/{assessment.id}/submit/", {"competency": plan.competency_id, "question_ids": [question.id for question in questions], "answers": first_answers}, format="json")
        self.assertEqual(submitted.status_code, 201, submitted.data)
        self.assertEqual(submitted.data["incorrect_question_ids"], [questions[1].id])

        retried = self.client.get(f"/api/assessments/{assessment.id}/start/?competency={plan.competency_id}&questions={questions[1].id}")
        self.assertEqual(retried.status_code, 200, retried.data)
        self.assertEqual([item["id"] for item in retried.data["assessment"]["questions"]], [questions[1].id])
        passed = self.client.post(f"/api/assessments/{assessment.id}/submit/", {"competency": plan.competency_id, "question_ids": [questions[1].id], "answers": [{"question_id": questions[1].id, "answer": questions[1].correct_answer}]}, format="json")
        self.assertEqual(passed.status_code, 201, passed.data)
        plan.refresh_from_db()
        self.assertEqual(plan.status, "completed")

    @patch("recovery.views.get_llm_provider")
    def test_teacher_can_generate_editable_questions_for_selected_competencies(self, provider_factory):
        self.authenticate("teacher@tala.edu.ph")
        assessment = Assessment.objects.filter(subject__in=get_user_model().objects.get(username="teacher@tala.edu.ph").tala_profile.assigned_subjects.all()).first()
        assessment.is_active = False
        assessment.save(update_fields=["is_active"])
        competency = Competency.objects.filter(subject=assessment.subject, is_active=True).first()
        provider_factory.return_value.generate.return_value = LLMResponse(
            text=json.dumps({"questions": [{"competency_code": competency.code, "prompt": "Which value satisfies x + 2 = 5?", "question_type": "mcq", "options": ["1", "2", "3", "4"], "correct_answer": "3"}]}),
            provider="test-provider", model="test-model",
        )

        response = self.client.post(f"/api/assessments/{assessment.id}/generate-questions/", {"competency_ids": [competency.id], "count": 1, "question_type": "mcq"}, format="json")

        self.assertEqual(response.status_code, 201, response.data)
        self.assertEqual(len(response.data["questions"]), 1)
        generated = Question.objects.get(pk=response.data["questions"][0]["id"])
        self.assertEqual(generated.assessment, assessment)
        self.assertEqual(generated.competency, competency)
        self.assertFalse(assessment.is_active)

    def test_teacher_can_delete_unused_draft_assessment_but_not_learner_records(self):
        self.authenticate("teacher@tala.edu.ph")
        teacher = get_user_model().objects.get(username="teacher@tala.edu.ph")
        subject = teacher.tala_profile.assigned_subjects.first()
        draft = Assessment.objects.create(title="Unused teacher draft", subject=subject, kind=Assessment.Kind.PRE, created_by=teacher, is_active=False)

        deleted = self.client.delete(f"/api/assessments/{draft.id}/")

        self.assertEqual(deleted.status_code, 204, deleted.data)
        self.assertFalse(Assessment.objects.filter(pk=draft.id).exists())
        retained = Assessment.objects.create(title="Draft with learner record", subject=subject, kind=Assessment.Kind.PRE, created_by=teacher, is_active=False)
        AssessmentAttempt.objects.create(assessment=retained, student=get_user_model().objects.get(username="student@tala.edu.ph"))
        blocked = self.client.delete(f"/api/assessments/{retained.id}/")
        self.assertEqual(blocked.status_code, 409, blocked.data)
        self.assertTrue(Assessment.objects.filter(pk=retained.id).exists())

    def test_teacher_can_delete_individual_draft_question_but_not_active_question(self):
        self.authenticate("teacher@tala.edu.ph")
        teacher = get_user_model().objects.get(username="teacher@tala.edu.ph")
        subject = teacher.tala_profile.assigned_subjects.first()
        competency = Competency.objects.filter(subject=subject).first()
        draft = Assessment.objects.create(title="Question deletion draft", subject=subject, kind=Assessment.Kind.PRE, created_by=teacher, is_active=False)
        question = Question.objects.create(assessment=draft, competency=competency, prompt="Remove this draft question?", question_type=Question.QuestionType.TRUE_FALSE, options=["True", "False"], correct_answer="True")

        deleted = self.client.delete(f"/api/assessments/{draft.id}/questions/{question.id}/")

        self.assertEqual(deleted.status_code, 204, deleted.data)
        self.assertFalse(Question.objects.filter(pk=question.id).exists())
        active = Assessment.objects.create(title="Active question protection", subject=subject, kind=Assessment.Kind.PRE, created_by=teacher, is_active=True)
        protected = Question.objects.create(assessment=active, competency=competency, prompt="Keep this active question.", question_type=Question.QuestionType.TRUE_FALSE, options=["True", "False"], correct_answer="True")
        blocked = self.client.delete(f"/api/assessments/{active.id}/questions/{protected.id}/")
        self.assertEqual(blocked.status_code, 409, blocked.data)
        self.assertTrue(Question.objects.filter(pk=protected.id).exists())

    def test_remedial_exam_requires_parent_consent_without_locking_mastery(self):
        nina = get_user_model().objects.get(username="student@tala.edu.ph")
        RecoveryActivity.objects.filter(plan__student=nina).update(completed_at=timezone.now())
        remedial = Assessment.objects.get(kind=Assessment.Kind.REMEDIAL)
        self.authenticate("admin@tala.edu.ph")
        requested = self.client.post(f"/api/assessments/{remedial.id}/request-consent/", {"student": nina.id, "guardian_id": nina.tala_profile.guardian_contacts.get().id}, format="json")
        self.assertEqual(requested.status_code, 201, requested.data)
        consent = RemedialExamConsent.objects.get(pk=requested.data["id"])
        self.authenticate("student@tala.edu.ph")
        before = self.client.get("/api/assessments/")
        serialized = next(item for item in before.data["results"] if item["id"] == remedial.id)
        self.assertFalse(serialized["available"])
        self.assertEqual(serialized["consent_status"], "requested")
        mastery = next(item for item in before.data["results"] if item["kind"] == Assessment.Kind.POST)
        self.assertTrue(mastery["available"])
        self.client.credentials()
        token = signing.dumps({"consent_id": consent.id}, salt="tala-remedial-consent", compress=True)
        signed_form = SimpleUploadedFile("signed-consent.pdf", b"%PDF-1.4 signed consent", content_type="application/pdf")
        approved = self.client.post(f"/api/remedial-consent/?token={token}", {"decision": "approved", "signed_name": consent.guardian_name, "evidence_file": signed_form}, format="multipart")
        self.assertEqual(approved.status_code, 200, approved.data)
        self.assertTrue(approved.data["evidence_present"])
        consent.refresh_from_db()
        self.assertTrue(bool(consent.evidence_file))
        self.authenticate("student@tala.edu.ph")
        after = self.client.get("/api/assessments/")
        serialized = next(item for item in after.data["results"] if item["id"] == remedial.id)
        self.assertTrue(serialized["available"])
        self.assertEqual(serialized["consent_status"], "approved")

    @patch("recovery.tutor_views.get_llm_provider")
    def test_teacher_insight_is_evidence_bounded_and_class_scoped(self, provider_factory):
        provider_factory.return_value.generate.return_value = LLMResponse(text="Prioritize unlike denominators [E1] and review the active plan [P1].", provider="test", model="test-model")
        student = get_user_model().objects.get(username="student@tala.edu.ph")
        competency = RecoveryPlan.objects.filter(student=student).first().competency
        LearnerCompetencyEvidence.objects.create(student=student, competency=competency, evidence_type=LearnerCompetencyEvidence.EvidenceType.INTERVENTION, source_type="test", source_id=999999, score=60, summary="Needs guided practice with denominators.")
        self.authenticate("teacher@tala.edu.ph")

        response = self.client.post(f"/api/tutor/learners/{student.id}/insight/", {}, format="json")

        self.assertEqual(response.status_code, 200, response.data)
        self.assertEqual(response.data["provider"], "test")
        self.assertGreater(response.data["evidence_count"], 0)
        request = provider_factory.return_value.generate.call_args.args[0]
        self.assertIn("Cite supplied identifiers", request.system)
        self.assertIn("Learner evidence", request.messages[0]["content"])

    def test_approved_remedial_consent_cannot_be_revoked_after_exam_starts(self):
        nina = get_user_model().objects.get(username="student@tala.edu.ph")
        RecoveryActivity.objects.filter(plan__student=nina).update(completed_at=timezone.now())
        administrator = get_user_model().objects.get(username="admin@tala.edu.ph")
        remedial = Assessment.objects.get(kind=Assessment.Kind.REMEDIAL)
        consent = RemedialExamConsent.objects.create(
            assessment=remedial,
            student=nina,
            guardian=nina.tala_profile.guardian_contacts.get(),
            guardian_name="Rosa Flores",
            guardian_relationship="Parent",
            guardian_email="rosa.flores@example.com",
            status=RemedialExamConsent.Status.APPROVED,
            consent_text="Approved for this remedial exam.",
            requested_by=administrator,
        )
        self.authenticate("student@tala.edu.ph")
        started = self.client.get(f"/api/assessments/{remedial.id}/start/")
        self.assertEqual(started.status_code, 200, started.data)
        self.client.credentials()
        token = signing.dumps({"consent_id": consent.id}, salt="tala-remedial-consent", compress=True)

        revoked = self.client.post(f"/api/remedial-consent/?token={token}", {"decision": "revoked", "signed_name": consent.guardian_name}, format="json")

        self.assertEqual(revoked.status_code, 409, revoked.data)
        consent.refresh_from_db()
        self.assertEqual(consent.status, RemedialExamConsent.Status.APPROVED)

    def test_teacher_can_import_review_and_publish_docx_exam(self):
        self.authenticate("teacher@tala.edu.ph")
        from .models import AcademicClass, Competency, Subject
        subject = Subject.objects.get(code="GM")
        competency = Competency.objects.get(code="GM-02")
        academic_class = AcademicClass.objects.get(name="Rizal")
        document = Document()
        for line in ["1. What is the least common denominator of 3 and 4?", "A. 7", "B. 8", "C. 12", "D. 24", "Answer: C", "Competency: GM-02"]:
            document.add_paragraph(line)
        buffer = BytesIO(); document.save(buffer)
        upload = SimpleUploadedFile("fractions-exam.docx", buffer.getvalue(), content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document")
        with tempfile.TemporaryDirectory() as media_root, override_settings(MEDIA_ROOT=media_root):
            created = self.client.post("/api/content-imports/", {"title": "Imported fractions exam", "kind": "exam", "subject": subject.id, "competency": competency.id, "configuration": json.dumps({"assessment_kind": "pre", "assigned_class_ids": [academic_class.id]}), "source_file": upload}, format="multipart")
            self.assertEqual(created.status_code, 201, created.data)
            self.assertEqual(created.data["status"], ContentImport.Status.NEEDS_REVIEW)
            self.assertEqual(created.data["extracted_payload"]["question_count"], 1)
            self.assertTrue(Notification.objects.filter(recipient__username="admin@tala.edu.ph", kind=Notification.Kind.CONTENT_REVIEW, action_url=f"/imports/{created.data['id']}").exists())
            self.authenticate("admin@tala.edu.ph")
            published = self.client.post(f"/api/content-imports/{created.data['id']}/publish/", {}, format="json")
            self.assertEqual(published.status_code, 200, published.data)
            assessment = Assessment.objects.get(title="Imported fractions exam")
            self.assertFalse(assessment.is_active)
            self.assertEqual(assessment.questions.get().correct_answer, "12")

    def test_module_quiz_is_extracted_reviewed_and_required_for_completion(self):
        self.authenticate("teacher@tala.edu.ph")
        subject = Subject.objects.get(code="GM")
        competency = Competency.objects.get(code="GM-03")
        academic_class = AcademicClass.objects.get(name="Rizal")
        document = Document()
        document.add_paragraph("Review adding fractions with unlike denominators before answering the module check.")
        for line in ["1. What is 1/2 + 1/4?", "A. 2/6", "B. 3/4", "C. 1/8", "Answer: B"]:
            document.add_paragraph(line)
        buffer = BytesIO(); document.save(buffer)
        upload = SimpleUploadedFile("fraction-module.docx", buffer.getvalue(), content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document")
        with tempfile.TemporaryDirectory() as media_root, override_settings(MEDIA_ROOT=media_root):
            created = self.client.post("/api/content-imports/", {"title": "Fractions module with quiz", "kind": "module", "subject": subject.id, "competency": competency.id, "configuration": json.dumps({"assigned_class_ids": [academic_class.id]}), "source_file": upload}, format="multipart")
            self.assertEqual(created.status_code, 201, created.data)
            self.assertEqual(created.data["extracted_payload"]["practice_question_count"], 1)
            self.authenticate("admin@tala.edu.ph")
            published = self.client.post(f"/api/content-imports/{created.data['id']}/publish/", {}, format="json")
            self.assertEqual(published.status_code, 200, published.data)
            assignment = LearningAssignment.objects.get(resource__title="Fractions module with quiz")
            question = assignment.resource.practice_questions.get()
            self.authenticate()
            blocked = self.client.post(f"/api/learning-assignments/{assignment.id}/complete/", {}, format="json")
            self.assertEqual(blocked.status_code, 409)
            submitted = self.client.post(f"/api/learning-assignments/{assignment.id}/submit-quiz/", {"answers": [{"question_id": question.id, "answer": question.correct_answer}]}, format="json")
            self.assertEqual(submitted.status_code, 201, submitted.data)
            self.assertTrue(submitted.data["passed"])
            self.assertIsNotNone(submitted.data["assignment"]["completed_at"])
            self.authenticate("teacher@tala.edu.ph")
            analytics = self.client.get("/api/dashboard/teacher/materials/")
            self.assertEqual(analytics.status_code, 200, analytics.data)
            material = next(item for item in analytics.data["materials"] if item["id"] == assignment.id)
            self.assertEqual(material["completed"], 1)
            self.assertEqual(material["average_quiz_score"], 100.0)
            revised_payload = created.data["extracted_payload"]
            revised_payload["practice_questions"][0]["correct_answer"] = "2/6"
            revised = self.client.patch(f"/api/content-imports/{created.data['id']}/", {"extracted_payload": revised_payload}, format="json")
            self.assertEqual(revised.status_code, 200, revised.data)
            assignment.resource.refresh_from_db()
            self.assertEqual(assignment.resource.practice_questions.get().correct_answer, "2/6")
            revised_payload["practice_questions"] = []
            revised_payload["practice_question_count"] = 0
            removed = self.client.patch(f"/api/content-imports/{created.data['id']}/", {"extracted_payload": revised_payload}, format="json")
            self.assertEqual(removed.status_code, 200, removed.data)
            assignment.resource.refresh_from_db()
            self.assertFalse(assignment.resource.practice_questions.exists())

    def test_teacher_can_publish_video_learning_material(self):
        self.authenticate("teacher@tala.edu.ph")
        from .models import Competency, Subject
        subject = Subject.objects.get(code="GM")
        competency = Competency.objects.get(code="GM-03")
        upload = SimpleUploadedFile("fraction-demo.mp4", b"small-test-video", content_type="video/mp4")
        with tempfile.TemporaryDirectory() as media_root, override_settings(MEDIA_ROOT=media_root):
            academic_class = AcademicClass.objects.get(name="Rizal")
            created = self.client.post("/api/content-imports/", {"title": "Fraction demonstration", "kind": "video", "subject": subject.id, "competency": competency.id, "configuration": json.dumps({"description": "Teacher demonstration of adding unlike fractions.", "assigned_class_ids": [academic_class.id]}), "source_file": upload}, format="multipart")
            self.assertEqual(created.status_code, 201, created.data)
            self.assertEqual(created.data["uploaded_by_name"], "Elena Cruz")
            self.authenticate("ramon.mendoza@tala.edu.ph")
            other_teacher_imports = self.client.get("/api/content-imports/")
            self.assertFalse(any(item["id"] == created.data["id"] for item in other_teacher_imports.data["results"]))
            self.authenticate("admin@tala.edu.ph")
            published = self.client.post(f"/api/content-imports/{created.data['id']}/publish/", {}, format="json")
            self.assertEqual(published.status_code, 200, published.data)
            resource = LearningResource.objects.get(title="Fraction demonstration")
            self.assertEqual(resource.resource_type, LearningResource.ResourceType.VIDEO)
            self.assertTrue(resource.is_approved)
            assignment = LearningAssignment.objects.get(resource=resource)
            self.assertEqual(list(assignment.assigned_classes.values_list("id", flat=True)), [academic_class.id])
            self.authenticate()
            assigned = self.client.get("/api/learning-assignments/")
            self.assertEqual(assigned.status_code, 200)
            assigned_item = next(item for item in assigned.data["results"] if item["id"] == assignment.id)
            self.assertIsNone(assigned_item["completed_at"])
            opened = self.client.post(f"/api/learning-assignments/{assignment.id}/open/", {}, format="json")
            self.assertEqual(opened.status_code, 200)
            self.assertIsNotNone(opened.data["opened_at"])
            playback = self.client.post(f"/api/learning-assignments/{assignment.id}/progress/", {"position_seconds": 48, "duration_seconds": 120}, format="json")
            self.assertEqual(playback.status_code, 200)
            self.assertEqual(playback.data["playback_position_seconds"], 48)
            self.assertEqual(playback.data["progress_percent"], 40)
            blocked = self.client.post(f"/api/learning-assignments/{assignment.id}/complete/", {}, format="json")
            self.assertEqual(blocked.status_code, 409)
            playback = self.client.post(f"/api/learning-assignments/{assignment.id}/progress/", {"position_seconds": 108, "duration_seconds": 120}, format="json")
            self.assertEqual(playback.data["progress_percent"], 90)
            completed = self.client.post(f"/api/learning-assignments/{assignment.id}/complete/", {}, format="json")
            self.assertEqual(completed.status_code, 200)
            self.assertIsNotNone(completed.data["completed_at"])
            self.authenticate("admin@tala.edu.ph")
            detail = self.client.get(f"/api/resources/{resource.id}/")
            self.assertEqual(detail.status_code, 200)
            signed_url = urlsplit(detail.data["file_url"])
            self.client.credentials()
            protected_file = self.client.get(f"{signed_url.path}?{signed_url.query}")
            self.assertEqual(protected_file.status_code, 200)
            self.assertNotIn("X-Frame-Options", protected_file)
            self.authenticate("admin@tala.edu.ph")
            archived = self.client.post(f"/api/content-imports/{created.data['id']}/archive/", {}, format="json")
            self.assertEqual(archived.status_code, 200, archived.data)
            self.assertIsNotNone(archived.data["archived_at"])
            resource.refresh_from_db(); assignment.refresh_from_db()
            self.assertFalse(resource.is_approved)
            self.assertFalse(assignment.is_active)

    @override_settings(WHISPER_ENABLED=True)
    @patch("recovery.content_imports.transcribe_video")
    def test_local_video_transcript_can_extract_quiz_questions(self, transcribe_video):
        transcribe_video.return_value = """Video lesson transcript.
1. Which practice keeps an account secure?
A. Share the password
B. Use a strong password
C. Disable updates
D. Post credentials
Answer: B
"""
        self.authenticate("teacher@tala.edu.ph")
        subject = Subject.objects.get(code="GM")
        competency = Competency.objects.get(code="GM-03")
        upload = SimpleUploadedFile("secure-video.mp4", b"test-video", content_type="video/mp4")
        with tempfile.TemporaryDirectory() as media_root, override_settings(MEDIA_ROOT=media_root):
            created = self.client.post("/api/content-imports/", {"title": "Video with spoken quiz", "kind": "video", "subject": subject.id, "competency": competency.id, "configuration": json.dumps({"assigned_class_ids": [AcademicClass.objects.get(name="Rizal").id]}), "source_file": upload}, format="multipart")
            self.assertEqual(created.status_code, 201, created.data)
            self.assertEqual(created.data["extracted_payload"]["transcription_status"], "completed")
            self.assertEqual(created.data["extracted_payload"]["practice_question_count"], 1)
            self.authenticate("admin@tala.edu.ph")
            published = self.client.post(f"/api/content-imports/{created.data['id']}/publish/", {}, format="json")
            self.assertEqual(published.status_code, 200, published.data)
            self.assertEqual(LearningResource.objects.get(pk=published.data["published_resource"]).practice_questions.count(), 1)

    @patch("recovery.content_imports.get_llm_provider")
    def test_video_transcript_ai_quiz_is_grounded_and_reviewable(self, provider_factory):
        competency = Competency.objects.get(code="GM-03")
        rows = [
            {
                "prompt": f"Grounded transcript question {index}?",
                "options": ["A", "B", "C", "D"],
                "correct_answer": "B",
                "explanation": "The transcript supports B.",
                "source_locator": f"Transcript topic {index}",
            }
            for index in range(1, 4)
        ]
        provider_factory.return_value.generate.return_value = LLMResponse(
            text=json.dumps({"questions": rows}), provider="test", model="grounded-model"
        )

        questions, response = generate_grounded_practice_questions("A sufficiently clear lesson transcript.", competency)

        self.assertEqual(len(questions), 3)
        self.assertTrue(all(item["provenance"] == "ai" for item in questions))
        self.assertTrue(all(item["confidence"] == "needs_review" for item in questions))
        self.assertEqual(response.model, "grounded-model")

    @patch("recovery.content_imports.get_llm_provider")
    def test_video_quiz_accepts_wrapped_json_and_letter_answers(self, provider_factory):
        competency = Competency.objects.get(code="GM-03")
        rows = [
            {
                "question": f"Grounded wrapped question {index}?",
                "choices": ["First choice", "Correct choice", "Third choice", "Fourth choice"],
                "answer": "B",
                "explanation": "The transcript identifies the second choice.",
            }
            for index in range(1, 4)
        ]
        provider_factory.return_value.generate.return_value = LLMResponse(
            text=f"Here is the requested quiz:\n```json\n{json.dumps({'questions': rows})}\n```\nPlease review it.",
            provider="test",
            model="grounded-model",
        )

        questions, _ = generate_grounded_practice_questions("A sufficiently clear lesson transcript.", competency)

        self.assertEqual(len(questions), 3)
        self.assertTrue(all(item["correct_answer"] == "Correct choice" for item in questions))

    def test_protected_media_rejects_anonymous_unsigned_requests(self):
        with tempfile.TemporaryDirectory() as media_root, override_settings(MEDIA_ROOT=media_root):
            resource = LearningResource.objects.create(
                title="Protected module",
                resource_type=LearningResource.ResourceType.MODULE,
                file=SimpleUploadedFile("protected.pdf", b"private-content", content_type="application/pdf"),
                original_filename="protected.pdf",
                mime_type="application/pdf",
                is_approved=True,
            )
            response = self.client.get(f"/api/resources/{resource.id}/file/")
            self.assertEqual(response.status_code, 403)

    def test_protected_video_supports_mobile_byte_range_requests(self):
        with tempfile.TemporaryDirectory() as media_root, override_settings(MEDIA_ROOT=media_root, DEBUG=True):
            resource = LearningResource.objects.create(
                title="Mobile video",
                resource_type=LearningResource.ResourceType.VIDEO,
                file=SimpleUploadedFile("mobile-video.mp4", b"0123456789", content_type="video/mp4"),
                original_filename="mobile-video.mp4",
                mime_type="video/mp4",
                is_approved=True,
            )
            self.authenticate("admin@tala.edu.ph")
            detail = self.client.get(f"/api/resources/{resource.id}/")
            signed_url = urlsplit(detail.data["file_url"])
            self.client.credentials()

            response = self.client.get(f"{signed_url.path}?{signed_url.query}", HTTP_RANGE="bytes=2-5")

            self.assertEqual(response.status_code, 206)
            self.assertEqual(response["Accept-Ranges"], "bytes")
            self.assertEqual(response["Content-Range"], "bytes 2-5/10")
            self.assertEqual(response["Content-Length"], "4")
            self.assertEqual(b"".join(response.streaming_content), b"2345")

    def test_notifications_are_scoped_to_current_user(self):
        self.authenticate()
        response = self.client.get("/api/notifications/")
        self.assertEqual(response.status_code, 200)
        self.assertGreaterEqual(len(response.data["results"]), 1)
        notification_id = response.data["results"][0]["id"]
        marked = self.client.post(f"/api/notifications/{notification_id}/read/")
        self.assertEqual(marked.status_code, 200)
        self.assertTrue(marked.data["is_read"])
